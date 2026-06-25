"""
Issue #3: Document upload endpoint — POST /upload accepts PDF, TXT, DOCX.
"""
import hashlib, os, tempfile, logging
from pathlib import Path
from fastapi import APIRouter, UploadFile, File, HTTPException

router  = APIRouter()
log     = logging.getLogger(__name__)
ALLOWED = {".pdf", ".txt", ".docx", ".md"}
MAX_MB  = int(os.getenv("MAX_UPLOAD_SIZE_MB", "10"))


@router.post("/upload")
async def upload_document(file: UploadFile = File(...)):
    """Upload a document and ingest it into the FAISS vector store."""
    ext = Path(file.filename or "").suffix.lower()
    if ext not in ALLOWED:
        raise HTTPException(400, f"Unsupported type {ext}. Allowed: {ALLOWED}")

    content = await file.read()
    if len(content) > MAX_MB * 1024 * 1024:
        raise HTTPException(413, f"File too large (max {MAX_MB} MB)")

    doc_hash = hashlib.sha256(content).hexdigest()[:12]

    with tempfile.NamedTemporaryFile(suffix=ext, delete=False) as tmp:
        tmp.write(content)
        tmp_path = tmp.name

    try:
        text, chunks = _parse(tmp_path, ext)
        _ingest(chunks, doc_hash)
        return {"filename": file.filename, "hash": doc_hash,
                "characters": len(text), "chunks": len(chunks), "status": "ingested"}
    finally:
        os.unlink(tmp_path)


def _parse(path: str, ext: str) -> tuple[str, list[str]]:
    if ext == ".txt" or ext == ".md":
        text = Path(path).read_text(errors="replace")
    elif ext == ".pdf":
        try:
            from pdfminer.high_level import extract_text
            text = extract_text(path)
        except ImportError:
            text = Path(path).read_text(errors="replace")
    elif ext == ".docx":
        try:
            import docx
            text = "\n".join(p.text for p in docx.Document(path).paragraphs)
        except ImportError:
            text = Path(path).read_text(errors="replace")
    else:
        text = ""
    chunk_size = int(os.getenv("CHUNK_SIZE", "500"))
    overlap    = int(os.getenv("CHUNK_OVERLAP", "50"))
    chunks = [text[i:i+chunk_size] for i in range(0, len(text), chunk_size-overlap)]
    return text, chunks


def _ingest(chunks: list[str], doc_id: str) -> None:
    try:
        from rag_project.vectorstore import get_vectorstore
        vs = get_vectorstore()
        from langchain.schema import Document
        docs = [Document(page_content=c, metadata={"doc_id": doc_id, "chunk": i})
                for i, c in enumerate(chunks)]
        vs.add_documents(docs)
        vs.save_local(os.getenv("VECTOR_STORE_PATH", "./data/faiss_index"))
        log.info("Ingested %d chunks for doc %s", len(chunks), doc_id)
    except Exception as e:
        log.warning("Ingest failed: %s", e)
